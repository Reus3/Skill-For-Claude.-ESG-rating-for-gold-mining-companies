#!/usr/bin/env python3
"""
Генератор академического Word-отчёта по результатам ESG-расчёта (v2, с АКРА-элементами).

Новое в этой версии:
 - Упрощённые таблицы блоков: «Вес блока: 333, компания набрала: 114,0 (34%)»
 - Блоки «Лучшие практики», «Проблемные зоны», «Потенциал роста»
 - Ссылки на страницы PDF для каждой ненулевой строки
 - Матрица отраслевых рисков 10×4 (в стиле АКРА)
 - Индикатор качества анализа (coverage) и индикатор запаса по контексту
 - Счётчик токенов (если передан в input)
 - Санкционный бонус отображается отдельно

Использование:
    python build_docx_report.py result.json output.docx
    python build_docx_report.py result.json output.docx --meta meta.json
"""
import json
import sys
import argparse
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Требуется python-docx: pip install python-docx --break-system-packages", file=sys.stderr)
    sys.exit(1)


BLOCK_NAMES = {"E": "Экология", "S": "Социальная ответственность", "G": "Корпоративное управление"}

INDUSTRY_RISKS = [
    ("Недостаточная социальная ответственность бизнеса", "S7, S10"),
    ("Безопасность работников", "S1, S2, S3"),
    ("Здоровье работников", "S4, S6"),
    ("Профессиональные заболевания", "S4"),
    ("Нарушение прав человека, дискриминация", "S7"),
    ("Отсутствие / утечка квалифицированных кадров", "S5, S6"),
    ("Снижение лояльности сотрудников", "S5, S9"),
    ("Управление хвостохранилищами (TSF)", "E6"),
    ("Обращение с опасными веществами (цианиды, ртуть)", "E7"),
    ("Коррупционные риски", "G3, G4, G7"),
]

BEST_PRACTICES_LIBRARY = {
    "E1": {
        "text": "Newmont и Agnico Eagle раскрывают Scope 1+2 по методологии GHG Protocol в ESG databook с явно описанными границами консолидации (equity share vs operational control). Barrick публикует удельные tCO2e/oz AuEq с 3-летней динамикой.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024 / ESG Databook", "section": "Climate Change → GHG Emissions table",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Climate Change → Energy & Emissions Performance",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Climate Change → Carbon Footprint",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "E2": {
        "text": "Newmont раскрывает Scope 3 по категориям 1 (purchased goods), 4 (upstream transportation) и 11 (use of sold products) с методологией оценки. Barrick и Agnico Eagle проходят verification по ISO 14064-3.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Climate Change → Scope 3 Emissions Inventory",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Climate Change → GHG Inventory & Verification",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Climate Change → Scope 3 Disclosure",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "E3": {
        "text": "Agnico Eagle раскрывает структуру энергобаланса в ГДж с разбивкой по топливу (diesel, HFO), электричеству (сеть vs собственная генерация) и теплу. Удельный показатель приводится ежегодно в ESG databook.",
        "sources": [
            {"company": "Agnico Eagle", "document": "ESG Databook (Sustainability Report) for FY2024", "section": "Energy Performance Tables",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "E4": {
        "text": "Полюс поддерживает 100% ВИЭ в электропотреблении с 2021 года через российские зелёные сертификаты. Gold Fields применяет PPA с солнечными фермами в ЮАР и Чили. Agnico Eagle — гидроэнергетика в Квебеке и Финляндии.",
        "sources": [
            {"company": "ПАО «Полюс»", "document": "Отчёт об устойчивом развитии за 2024 год", "section": "Климат и энергетика → Возобновляемая энергия",
             "url": "https://polyus.com/ru/sustainability/"},
            {"company": "Gold Fields", "document": "Integrated Annual Report for FY2024", "section": "Climate Strategy → Renewable Energy & PPAs",
             "url": "https://www.goldfields.com/reports.php"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Climate Change → Renewable Power Sources",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "E5": {
        "text": "Newmont применяет WRI Aqueduct для идентификации операций в зонах водного стресса и публикует water stewardship programs. Gold Fields публикует water balance и water recycling rate по каждой площадке.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Water → Water Stewardship & Aqueduct mapping",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Gold Fields", "document": "Integrated Annual Report for FY2024", "section": "Water Management → Site-level Water Balance",
             "url": "https://www.goldfields.com/reports.php"},
        ],
    },
    "E6": {
        "text": "Newmont и Barrick подписали GISTM и назначили Accountable Executive + Engineer of Record по каждому TSF. Agnico Eagle публикует отдельный Tailings Facilities Disclosure по модели RMI.",
        "sources": [
            {"company": "Newmont", "document": "Tailings Management Disclosure for FY2024 (отдельный документ)",
             "section": "GISTM Conformance → Accountable Executive & EoR",
             "url": "https://www.newmont.com/sustainability/tailings-management"},
            {"company": "Barrick", "document": "Tailings Management Disclosure for FY2024",
             "section": "GISTM Implementation Status",
             "url": "https://www.barrick.com/English/sustainability/tailings"},
            {"company": "Agnico Eagle", "document": "Tailings Facilities Disclosure (RMI-aligned) for FY2024",
             "section": "Facility-by-facility risk classification",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "E7": {
        "text": "Newmont и Barrick полностью сертифицированы по ICMI Cyanide Code по всем применимым операциям. Agnico Eagle раскрывает объёмы потребления цианида и incident reporting.",
        "sources": [
            {"company": "Newmont", "document": "ICMI Cyanide Code Compliance Report for FY2024", "section": "Site Certification Status",
             "url": "https://cyanidecode.org/signatory/newmont/"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Cyanide Management → ICMI Recertification",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Hazardous Materials → Cyanide Use",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "E8": {
        "text": "Newmont имеет BAP (Biodiversity Action Plan) с no-net-loss обязательством. Barrick применяет biodiversity offsets. Gold Fields публикует отчёты по воздействию на ключевые виды.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Biodiversity & Land Use → Biodiversity Action Plans",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Biodiversity → Offsets & No-Net-Loss",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Gold Fields", "document": "Integrated Annual Report for FY2024", "section": "Environment → Biodiversity Performance",
             "url": "https://www.goldfields.com/reports.php"},
        ],
    },
    "E9": {
        "text": "Newmont раскрывает ARO-резервы по каждому активу с методологией оценки и сценарным анализом. Barrick обеспечивает полное покрытие банковскими гарантиями для крупных mine closure obligations.",
        "sources": [
            {"company": "Newmont", "document": "Annual Report 10-K for FY2024", "section": "Notes to Consolidated Financial Statements → Asset Retirement Obligations",
             "url": "https://www.newmont.com/investors/annual-report/"},
            {"company": "Barrick", "document": "Annual Information Form for FY2024", "section": "Mine Closure Obligations & Surety Bonds",
             "url": "https://www.barrick.com/English/investors/financial-reports"},
        ],
    },
    "E10": {
        "text": "Newmont, Barrick публикуют все экологические штрафы и significant spills (>1 barrel, >1 ton) с описанием причин и корректирующих мер.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Environmental Compliance → Fines & Significant Incidents",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Environmental Performance → Reportable Incidents",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "E11": {
        "text": "Newmont, Barrick, Agnico Eagle — сертификация ISO 14001 по 100% операционных активов с ежегодным surveillance audit.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Management Systems → ISO 14001 Coverage Map",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "EMS → Site Certifications",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Environmental Management → Certifications",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "S1": {
        "text": "Agnico Eagle — LTIFR на 1 млн часов с раскрытием методики (сотрудники + подрядчики, 1 млн часов base). Ежеквартальные обновления.",
        "sources": [
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Health & Safety Performance → TRIFR & LTIFR",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "S2": {
        "text": "Newmont — zero fatalities за 2024-2025 финансовые годы (лучший результат в отрасли). Barrick публикует FIFR и описание обстоятельств каждого fatality с корректирующими мерами.",
        "sources": [
            {"company": "Newmont", "document": "Annual Report 10-K for FY2024 (опубликован в 2025)", "section": "Operations Review → Health & Safety",
             "url": "https://www.newmont.com/investors/annual-report/"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Health & Safety → Fatality Prevention",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "S3": {
        "text": "Newmont, Barrick, Agnico Eagle — ISO 45001 по 100% операций. Полиметалл также достиг 100% охвата.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "OHS Management Systems → ISO 45001",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Health & Safety → ISO 45001 Certifications",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Health & Safety Management",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
            {"company": "АО «Полиметалл»", "document": "Годовой отчёт за 2024 год", "section": "Охрана труда и промышленная безопасность",
             "url": "https://polymetal.ru/"},
        ],
    },
    "S4": {
        "text": "Agnico Eagle публикует методологию учёта occupational diseases (по ILO criteria) с 5-летней ретроспективой.",
        "sources": [
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Health & Safety → Occupational Health metrics",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "S5": {
        "text": "Agnico Eagle раскрывает текучесть с разбивкой office / operational / rotational. Рыночный бенчмарк для вахты — 15-20% в год.",
        "sources": [
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Our People → Workforce Stability",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "S6": {
        "text": "Полиметалл за 2024 год — 132 часа обучения на сотрудника (лидер отрасли). Newmont — разбивка по категориям (safety / technical / leadership).",
        "sources": [
            {"company": "АО «Полиметалл»", "document": "Годовой отчёт за 2024 год", "section": "Забота о сотрудниках → Обучение и развитие",
             "url": "https://polymetal.ru/"},
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Our People → Training & Development table",
             "url": "https://www.newmont.com/sustainability"},
        ],
    },
    "S7": {
        "text": "Newmont, Barrick публикуют IBA-соглашения с коренными общинами (Canada, Australia) с раскрытием сумм выплат и грантов. Grievance mechanism соответствует UN Guiding Principles on Business and Human Rights.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Indigenous Peoples & Community Engagement → IBA Agreements",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Human Rights → Indigenous Peoples & Grievance Mechanism",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "S8": {
        "text": "Newmont — 100% критических поставщиков проходят ESG assessment по OECD Due Diligence Guidance for Responsible Supply Chains of Minerals. Barrick — публичный supplier code of conduct + audit коренных поставщиков.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Supply Chain → ESG Assessment Coverage",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Supplier Code of Conduct (revised 2024)", "section": "Sourcing Standards → Audit Process",
             "url": "https://www.barrick.com/English/sustainability/supply-chain/"},
        ],
    },
    "S9": {
        "text": "Agnico Eagle публикует число обращений по каналам и % урегулированных с разбивкой по типам (compensation, workplace conditions, safety).",
        "sources": [
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Our People → Grievance Mechanism Statistics",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "S10": {
        "text": "Newmont — community investment plans по каждой операции с измеримыми KPI (education, health, infrastructure). Полюс публикует разбивку социальных инвестиций по регионам и направлениям.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Communities → Community Investment programs",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "ПАО «Полюс»", "document": "Отчёт об устойчивом развитии за 2024 год", "section": "Социальные инвестиции → Разбивка по регионам",
             "url": "https://polyus.com/ru/sustainability/"},
        ],
    },
    "G1": {
        "text": "Newmont, Barrick — 8-9 из 10 директоров независимые, с раскрытыми критериями независимости (NYSE / TSX standards).",
        "sources": [
            {"company": "Newmont", "document": "Proxy Statement for FY2024", "section": "Board of Directors → Director Independence",
             "url": "https://www.newmont.com/investors/proxy-information/"},
            {"company": "Barrick", "document": "Management Information Circular for FY2024", "section": "Governance → Board Composition",
             "url": "https://www.barrick.com/English/investors/proxy/"},
        ],
    },
    "G2": {
        "text": "Newmont, Agnico Eagle — строгое разделение Chairman и CEO. Barrick применяет Lead Independent Director при combined role.",
        "sources": [
            {"company": "Newmont", "document": "Proxy Statement for FY2024", "section": "Board Leadership Structure",
             "url": "https://www.newmont.com/investors/proxy-information/"},
            {"company": "Agnico Eagle", "document": "Management Information Circular for FY2024", "section": "Board Leadership",
             "url": "https://www.agnicoeagle.com/English/investors"},
            {"company": "Barrick", "document": "Management Information Circular for FY2024", "section": "Lead Independent Director Role",
             "url": "https://www.barrick.com/English/investors/proxy/"},
        ],
    },
    "G3": {
        "text": "Newmont, Barrick — полная ISO 37001 сертификация (anti-bribery management system). Agnico Eagle — ABC framework с годовыми training KPI.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Ethics & Compliance → Anti-Corruption",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Code of Business Conduct & Ethics (revised 2024)", "section": "Anti-Bribery Standards",
             "url": "https://www.barrick.com/English/about/governance/ethics/"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Ethics & Compliance → ABC Framework",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
    "G4": {
        "text": "Newmont, Barrick публикуют число disciplinary dismissals за коррупцию и описывают процедуру расследований.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Ethics & Compliance → Anti-Corruption Cases",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Compliance → Confirmed Incidents",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "G5": {
        "text": "Agnico Eagle, Newmont — ESG-KPI интегрированы в short-term incentive plans руководства (health & safety, climate targets) с явными весами (~20-30% от STIP).",
        "sources": [
            {"company": "Agnico Eagle", "document": "Management Information Circular for FY2024", "section": "Executive Compensation → STIP ESG Modifiers",
             "url": "https://www.agnicoeagle.com/English/investors"},
            {"company": "Newmont", "document": "Proxy Statement for FY2024", "section": "Compensation Discussion & Analysis → ESG Performance Metrics",
             "url": "https://www.newmont.com/investors/proxy-information/"},
        ],
    },
    "G6": {
        "text": "Newmont — интегрированный enterprise risk management с quantitative risk appetite по climate и TSF рискам, отчётность Board of Directors ежеквартально.",
        "sources": [
            {"company": "Newmont", "document": "Annual Report 10-K for FY2024", "section": "Risk Management → Enterprise Risk Framework",
             "url": "https://www.newmont.com/investors/annual-report/"},
        ],
    },
    "G7": {
        "text": "Newmont, Barrick — whistleblower channels на 10+ языках с публикацией % урегулированных обращений и времени разрешения.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Ethics & Compliance → Speak-Up Line statistics",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Barrick", "document": "Sustainability Report for FY2024", "section": "Ethics & Compliance → Reporting Mechanisms",
             "url": "https://www.barrick.com/English/sustainability"},
        ],
    },
    "G8": {
        "text": "Newmont, Barrick, Agnico Eagle — полное членство ICMM + ежегодный LBMA RGG report + WGC RGMP self-assessment с external assurance. AngloGold Ashanti — RJC CoC certification.",
        "sources": [
            {"company": "Newmont", "document": "WGC RGMP Self-Assessment for FY2024", "section": "Independent Assurance Statement",
             "url": "https://www.gold.org/about-gold/responsible-gold/responsible-mining/principles/"},
            {"company": "Barrick", "document": "WGC RGMP Self-Assessment for FY2024", "section": "Conformance & Assurance",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Agnico Eagle", "document": "WGC RGMP Self-Assessment for FY2024", "section": "Conformance",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
            {"company": "AngloGold Ashanti", "document": "Sustainability Report for FY2024", "section": "Responsible Sourcing → RJC CoC Certification",
             "url": "https://www.anglogoldashanti.com/sustainability/"},
        ],
    },
    "G9": {
        "text": "Newmont, Barrick — JORC + NI 43-101 по всем активам, подписанные Competent Person / Qualified Person, с annual Mineral Reserve Statement + assurance.",
        "sources": [
            {"company": "Newmont", "document": "Annual Mineral Reserves & Resources Report for FY2024", "section": "Reserves & Resources by Asset → QP Sign-off",
             "url": "https://www.newmont.com/operations-and-projects/reserves-and-resources/"},
            {"company": "Barrick", "document": "Mineral Reserves & Resources Report for FY2024", "section": "Compliance with NI 43-101",
             "url": "https://www.barrick.com/English/operations/reserves-resources/"},
        ],
    },
    "G10": {
        "text": "Barrick публикует country-by-country tax report с налогами, роялти, payments to governments (EITI-aligned). Newmont — GRI 207 full disclosure.",
        "sources": [
            {"company": "Barrick", "document": "Tax Contribution Report (CbC) for FY2024", "section": "Country-by-Country Reporting Tables",
             "url": "https://www.barrick.com/English/sustainability"},
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "Economic Contribution → GRI 207 Disclosure",
             "url": "https://www.newmont.com/sustainability"},
        ],
    },
    "G11": {
        "text": "Newmont, Agnico Eagle — sustainability report с external limited assurance (Deloitte, KPMG), соответствие GRI + IFRS S1/S2 + SASB + TCFD.",
        "sources": [
            {"company": "Newmont", "document": "Sustainability Report for FY2024", "section": "About this Report → External Assurance Statement (KPMG)",
             "url": "https://www.newmont.com/sustainability"},
            {"company": "Agnico Eagle", "document": "Sustainability Report for FY2024", "section": "Reporting Frameworks → External Assurance (Deloitte)",
             "url": "https://www.agnicoeagle.com/English/sustainability"},
        ],
    },
}


def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_style_defaults(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(0)
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    for hname, size in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 13)]:
        h = doc.styles[hname]
        h.font.name = 'Times New Roman'
        h.font.size = Pt(size)
        h.font.color.rgb = RGBColor(0, 0, 0)
        h.font.bold = True
        h._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def add_kv_table(doc, rows, widths_cm=(6, 8)):
    t = doc.add_table(rows=len(rows), cols=2)
    t.style = 'Light Grid'
    for i, cm in enumerate(widths_cm):
        for cell in t.columns[i].cells:
            cell.width = Cm(cm)
    for i, (k, v) in enumerate(rows):
        t.cell(i, 0).text = str(k)
        t.cell(i, 1).text = str(v)
    doc.add_paragraph()
    return t


def add_multicol_table(doc, headers, rows, col_widths_cm=None, cell_fills=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid'
    for j, h in enumerate(headers):
        cell = t.cell(0, j)
        cell.text = h
        set_cell_background(cell, 'D5E8F0')
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            t.cell(i, j).text = str(val) if val is not None else ''
            if cell_fills and cell_fills[i - 1] and cell_fills[i - 1][j]:
                set_cell_background(t.cell(i, j), cell_fills[i - 1][j])
    if col_widths_cm:
        for j, cm in enumerate(col_widths_cm):
            for cell in t.columns[j].cells:
                cell.width = Cm(cm)
    doc.add_paragraph()
    return t


def write_title(doc, data):
    is_dynamic = "dynamic_analysis" in data
    ratings = data["ratings"]
    companies = list({r["company"] for r in ratings})
    years = sorted({r["year"] for r in ratings})

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Паспорт расчёта ESG-рейтинга")
    run.font.size = Pt(20); run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{', '.join(companies)} — {'–'.join(str(y) for y in years) if len(years) > 1 else years[0]}")
    run.font.size = Pt(14); run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Применённая методология: авторская disclosure-based, v3, шкала 0–999 баллов")
    run.font.italic = True; run.font.size = Pt(11)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Дата расчёта: {datetime.now().strftime('%d.%m.%Y')}")
    run.font.size = Pt(11)

    doc.add_paragraph()
    if not is_dynamic:
        r = ratings[0]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Итоговый балл: {r['total_score']:.1f} / 999")
        run.font.size = Pt(18); run.font.bold = True
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Рейтинг: {r['rating']['label']} — {r['rating']['interpretation']}")
        run.font.size = Pt(16); run.font.bold = True
    doc.add_page_break()


def write_object(doc, data):
    doc.add_heading("1. Объект оценки", level=1)
    for r in data["ratings"]:
        doc.add_heading(f"{r['company']}, {r['year']} год", level=2)
        add_kv_table(doc, [
            ("Название компании", r["company"]),
            ("Отчётный период", str(r["year"])),
            ("Применённая методология", "Авторская disclosure-based, v3 + санкционный корректив + матрица рисков"),
            ("Всего показателей", "32 (E=11, S=10, G=11)"),
            ("Максимальный балл", "999 (E=333, S=333, G=333)"),
        ])


def write_summary(doc, data):
    """Упрощённое резюме: вес блока — компания получила."""
    doc.add_heading("2. Резюме результатов", level=1)
    ratings = data["ratings"]

    if len(ratings) == 1:
        r = ratings[0]
        doc.add_paragraph("Таблица 1. Итоговый балл по блокам").runs[0].bold = True
        rows = []
        for block_code in ["E", "S", "G"]:
            bs = r["block_scores"][block_code]
            rows.append([
                f"Блок {block_code} — {BLOCK_NAMES[block_code]}",
                "333",
                f"{bs:.1f}",
                f"{(bs / 333 * 100):.0f}%",
            ])
        rows.append([
            "ИТОГО",
            "999",
            f"{r['total_score']:.1f}",
            f"{(r['total_score'] / 999 * 100):.0f}%",
        ])
        add_multicol_table(doc, ["Блок", "Вес блока", "Компания получила", "% от максимума"], rows,
                           col_widths_cm=[6.5, 3, 4, 3])

        p = doc.add_paragraph()
        p.add_run("Рейтинг: ").bold = True
        p.add_run(f"{r['rating']['label']} — {r['rating']['interpretation']}")

        if r.get("scaling_note"):
            p = doc.add_paragraph()
            p.add_run(f"Примечание: {r['scaling_note']}").italic = True
    else:
        doc.add_paragraph("Таблица 1. Динамика интегрального балла по годам").runs[0].bold = True
        headers = ["Год", "Балл", "Рейтинг", "Блок E", "Блок S", "Блок G", "ΔScore"]
        rows = []
        prev_score = None
        sorted_r = sorted(ratings, key=lambda x: x["year"])
        for r in sorted_r:
            delta = f"{r['total_score'] - prev_score:+.1f}" if prev_score is not None else "—"
            rows.append([
                r["year"], f"{r['total_score']:.1f}", r["rating"]["label"],
                f"{r['block_scores']['E']:.1f}", f"{r['block_scores']['S']:.1f}",
                f"{r['block_scores']['G']:.1f}", delta
            ])
            prev_score = r["total_score"]
        add_multicol_table(doc, headers, rows)


def write_meta_indicators(doc, data, meta):
    if not meta:
        return
    doc.add_heading("2.1. Индикаторы качества анализа (self-check)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Раздел показывает прозрачность работы skill: какой объём отчётности был проанализирован, "
        "насколько полно это удалось сделать, сколько ресурсов контекста было израсходовано "
        "и какой запас остаётся для дополнительных отчётов."
    ).italic = True

    tokens_used = meta.get("tokens_used", 0)
    context_budget = meta.get("context_budget", 200000)
    usage_pct = (tokens_used / context_budget * 100) if context_budget else 0

    if usage_pct < 30:
        context_status = "🟢 Низкая нагрузка: можно добавить ещё 2-3 отчёта сопоставимого размера без риска деградации качества."
    elif usage_pct < 60:
        context_status = "🟡 Средняя нагрузка: можно добавить ещё 1 отчёт аналогичного размера. При добавлении 2+ есть риск усечения цитат."
    else:
        context_status = "🔴 Высокая нагрузка: добавление нового отчёта с высокой вероятностью приведёт к потере деталей. Рекомендуется начать новую сессию."

    add_kv_table(doc, [
        ("Размер проанализированной отчётности (страниц PDF)", str(meta.get("total_pdf_pages", "—"))),
        ("Извлечено численных значений из отчётности", str(meta.get("values_extracted", "—"))),
        ("Токенов израсходовано на запуск skill", f"{tokens_used:,}".replace(",", " ") if tokens_used else "—"),
        ("Загрузка контекстного окна", f"{usage_pct:.1f}% (из {context_budget:,} токенов)".replace(",", " ") if tokens_used else "—"),
        ("Запас для добавления отчётов", context_status),
    ])

    # ---- Техническая проверка чтения PDF ----
    total_pages = meta.get("total_pdf_pages")
    pages_with_text = meta.get("pages_with_text")
    pages_no_text = meta.get("pages_without_text")
    ocr_warnings = meta.get("ocr_warnings", [])

    # Если переданы данные технической проверки — выводим отдельный подблок
    if total_pages is not None and pages_with_text is not None:
        doc.add_heading("2.1.1. Техническая проверка чтения PDF", level=2)
        p = doc.add_paragraph()
        p.add_run(
            "Раздел показывает, какие страницы исходного PDF удалось прочитать программно (извлечь "
            "текстовый слой), а какие могли остаться непрочитанными — например, из-за того, что они "
            "представлены сканами, изображениями или защищены от копирования. Эта проверка отвечает "
            "на вопрос «не пропустил ли skill часть данных по техническим причинам»."
        ).italic = True

        coverage_pdf = (pages_with_text / total_pages * 100) if total_pages else 0
        if coverage_pdf >= 95:
            pdf_status = "🟢 Полное покрытие: вся отчётность доступна для машинного чтения."
        elif coverage_pdf >= 80:
            pdf_status = ("🟡 Преимущественное покрытие: основной массив отчёта прочитан, "
                          "но отдельные страницы (вероятно, сканы или графические вставки) "
                          "могли остаться вне анализа.")
        else:
            pdf_status = ("🔴 Частичное покрытие: значительная часть отчёта не была прочитана "
                          "программно. Возможные причины: PDF состоит из сканов / отсутствует "
                          "текстовый слой / документ защищён. Рекомендуется применить OCR "
                          "(например, ABBYY FineReader, Tesseract) перед расчётом, либо "
                          "приложить дополнительный текстовый отчёт.")

        rows_pdf = [
            ("Всего страниц в PDF", str(total_pages)),
            ("Страниц с извлечённым текстовым слоем", f"{pages_with_text} ({coverage_pdf:.0f}%)"),
            ("Страниц без извлечённого текста (сканы / графика)", str(pages_no_text or 0)),
            ("Статус чтения PDF", pdf_status),
        ]
        add_kv_table(doc, rows_pdf)

        if ocr_warnings:
            p = doc.add_paragraph()
            p.add_run("⚠ Зафиксированные OCR-предупреждения: ").bold = True
            p.add_run("; ".join(ocr_warnings))

    # Индикатор качества верификации
    doc.add_heading("2.2. Качество верификации отчётности (coverage)", level=2)
    r = data["ratings"][0] if data["ratings"] else None
    if r:
        disclosed_count = sum(1 for i in r["indicators"] if i.get("x") is not None and i["x"] > 0)
        total_applicable = sum(1 for i in r["indicators"] if i.get("applicable", True))
        coverage_pct = disclosed_count / total_applicable * 100 if total_applicable else 0

        with_source = sum(1 for i in r["indicators"] if i.get("pdf_page_ref"))
        source_pct = with_source / disclosed_count * 100 if disclosed_count else 0

        if coverage_pct >= 80 and source_pct >= 80:
            quality_verdict = "🟢 Высокое качество верификации: ключевые данные из отчётности вряд ли упущены."
        elif coverage_pct >= 50:
            quality_verdict = "🟡 Среднее качество: часть показателей могла быть не найдена. Рекомендуется выборочная проверка аналитиком по показателям с x=0, имеющим высокий вес."
        else:
            quality_verdict = "🔴 Низкое качество: много нулевых показателей. Возможные причины: (а) сокращённый формат отчёта, (б) показатели раскрыты, но не найдены skill. Рекомендуется приложить отдельный ESG databook / отчёт об устойчивом развитии."

        add_kv_table(doc, [
            ("Покрытие показателей (x>0)", f"{disclosed_count} из {total_applicable} ({coverage_pct:.0f}%)"),
            ("Доля показателей со ссылкой на страницу PDF", f"{with_source} из {disclosed_count} ({source_pct:.0f}%)"),
            ("Итоговая оценка качества анализа", quality_verdict),
        ])

        high_weight_zeros = [i for i in r["indicators"]
                             if i.get("applicable", True) and (i.get("x") == 0 or i.get("x") is None)
                             and i["weight"] >= 25]
        if high_weight_zeros:
            p = doc.add_paragraph()
            p.add_run("Внимание: ").bold = True
            p.add_run(
                "следующие высоковесные показатели (вес ≥ 25) получили x=0. "
                "Это могут быть реально нераскрытые показатели (корректный результат) "
                "или пропущенные при автоматическом анализе. Рекомендуется выборочная проверка:"
            )
            for ind in sorted(high_weight_zeros, key=lambda x: -x["weight"])[:5]:
                doc.add_paragraph(f"• {ind['code']} ({ind['name']}) — вес {ind['weight']}", style='List Bullet')


def write_best_worst(doc, data):
    r = data["ratings"][-1] if data.get("dynamic_analysis") else data["ratings"][0]
    indicators = [i for i in r["indicators"] if i.get("applicable", True)]

    # --- 2.3 Лучшие практики ---
    doc.add_heading("2.3. Лучшие практики компании (top-5)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Показатели, по которым компания продемонстрировала сильное раскрытие и высокий балл. "
        "Это зоны, где компания соответствует или превосходит ожидания методологии."
    ).italic = True

    scored = [i for i in indicators if i.get("score") is not None and i["score"] > 0]
    best = sorted(scored, key=lambda i: -i.get("x", 0))[:5]
    best_rows = []
    for i in best:
        pct = int(i["x"] * 100) if i.get("x") else 0
        src = i.get("pdf_page_ref") or "—"
        best_rows.append([i["code"], i["name"], f"{i['score']:.1f} / {i['weight']}", f"{pct}%", src])
    if best_rows:
        add_multicol_table(doc,
            ["Код", "Показатель", "Балл", "% покрытия", "Источник в PDF"],
            best_rows, col_widths_cm=[1.5, 6, 2.5, 2, 4.4])
    else:
        doc.add_paragraph("Раскрытых показателей с ненулевым x не найдено.")

    # --- 2.4 Проблемные зоны ---
    doc.add_heading("2.4. Проблемные зоны (top-5 по уязвимости)", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Показатели, где компания получила низкое значение при высоком весе — раскрытые, но слабые практики. "
        "В отличие от раздела «Потенциал роста» здесь показаны показатели, по которым в отчёте есть данные, "
        "но их качество ниже средне-отраслевого уровня."
    ).italic = True
    problematic = [i for i in scored if i.get("x", 0) < 0.5]
    problematic.sort(key=lambda i: -(i["weight"] * (1 - i.get("x", 0))))
    prob_rows = []
    for i in problematic[:5]:
        pct = int(i["x"] * 100)
        lost = i["weight"] * (1 - i["x"])
        src = i.get("pdf_page_ref") or "—"
        caveat_info = ""
        if "zero_harm_cap" in i.get("caveats_applied", []):
            caveat_info = " [потолок fatalities]"
        elif "corruption_cap" in i.get("caveats_applied", []):
            caveat_info = " [потолок коррупция]"
        elif "active_community_conflict_cap" in i.get("caveats_applied", []):
            caveat_info = " [потолок конфликт]"
        prob_rows.append([i["code"], i["name"] + caveat_info, f"{i['score']:.1f} / {i['weight']}",
                          f"потеря {lost:.1f}", src])
    if prob_rows:
        add_multicol_table(doc,
            ["Код", "Показатель", "Балл", "Недобор", "Источник в PDF"],
            prob_rows, col_widths_cm=[1.5, 7, 2.5, 2.5, 2.9])
    else:
        doc.add_paragraph("Проблемных зон не выявлено — по всем раскрытым показателям x ≥ 0,5.")

    # --- 2.5 Потенциал роста (best practices library) ---
    doc.add_heading("2.5. Потенциал роста — примеры лучших практик в отрасли", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "Показатели, не раскрытые в отчёте (x = 0). Ниже приведены примеры того, "
        "как аналогичные показатели раскрывают лидеры отрасли — концентрированный "
        "референс по передовой практике для следующих отчётов компании."
    ).italic = True

    # Сводка использованных источников (для академической прозрачности)
    zero_inds = [i for i in indicators if i.get("x") == 0 or i.get("x") is None]
    zero_inds.sort(key=lambda i: -i["weight"])

    def _format_short(s):
        # Краткий формат для сводки в начале раздела: «Компания — документ»
        return f"{s['company']} — {s['document']}"

    def _format_full(s):
        # Полный формат под практикой: «Компания — документ, раздел X (URL)»
        parts = [f"{s['company']} — {s['document']}"]
        if s.get("section"):
            parts.append(f"раздел: {s['section']}")
        if s.get("url"):
            parts.append(f"источник: {s['url']}")
        return ", ".join(parts)

    # Собираем уникальный список источников по показателям, которые попадут в раздел
    all_sources_used = set()
    for ind in zero_inds[:6]:
        bp = BEST_PRACTICES_LIBRARY.get(ind["code"])
        if bp and isinstance(bp, dict):
            for src in bp.get("sources", []):
                all_sources_used.add(_format_short(src))

    if all_sources_used:
        p = doc.add_paragraph()
        p.add_run("Источники лучших практик в этом разделе: ").bold = True
        p.add_run("; ".join(sorted(all_sources_used)) + ".")
        # Пояснение про маркер «for FY2024» — для академической точности
        p_note = doc.add_paragraph()
        run = p_note.add_run(
            "Примечание о датировке: маркер «for FY2024» означает отчёт за финансовый 2024 год — "
            "он, как правило, публикуется в первом полугодии 2025 года. Маркер «for FY2024 (опубликован в 2025)» "
            "используется для документов, дата выхода которых формально позже отчётного периода. Указанные URL — "
            "корпоративные точки входа на разделы устойчивого развития / отчётности (актуальные документы могут "
            "архивироваться, но текущая версия всегда доступна по этим адресам)."
        )
        run.italic = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for ind in zero_inds[:6]:
        p = doc.add_paragraph()
        p.add_run(f"{ind['code']} — {ind['name']} (потенциал +{ind['weight']} баллов)").bold = True

        bp = BEST_PRACTICES_LIBRARY.get(ind["code"])
        if not bp:
            p = doc.add_paragraph()
            p.add_run("Как раскрывают лидеры отрасли: ").italic = True
            p.add_run("— (референс не занесён в справочник)")
            continue

        # Текст практики
        p = doc.add_paragraph()
        p.add_run("Как раскрывают лидеры отрасли: ").italic = True
        if isinstance(bp, dict):
            p.add_run(bp.get("text", ""))
            # Источники под каждой практикой — теперь с разделом и URL
            sources = bp.get("sources", [])
            if sources:
                p_src = doc.add_paragraph()
                run = p_src.add_run("Источники: ")
                run.italic = True
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                # Каждый источник на отдельной строке с переносом — так читабельнее
                for i, s in enumerate(sources):
                    sep = "." if i == len(sources) - 1 else ";"
                    run2 = p_src.add_run(_format_full(s) + sep + (" " if i < len(sources) - 1 else ""))
                    run2.italic = True
                    run2.font.size = Pt(10)
                    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        else:
            # Совместимость со старым плоским форматом
            p.add_run(str(bp))


def write_sanctions_adjustment(doc, data):
    r = data["ratings"][-1] if data.get("dynamic_analysis") else data["ratings"][0]

    # Отбираем все показатели, для которых санкционная информация была заявлена
    flagged = [i for i in r["indicators"] if i.get("sanctions_bonus")]
    if not flagged:
        return

    triggered = [i for i in flagged if i["sanctions_bonus"].get("triggered")]
    not_triggered = [i for i in flagged if not i["sanctions_bonus"].get("triggered")]

    doc.add_heading("3. Санкционный корректив", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Компания подвержена санкционным ограничениям, препятствующим членству в международных "
        "отраслевых организациях (ICMM, WGC, LBMA, RJC, EITI) или получению международно-признанных "
        "форм аудита (JORC, NI 43-101). Методология применяет санкционный бонус по формуле "
    )
    run = p.add_run("x_final = min(x_raw + 0,20; x_raw × 2; 1,0)")
    run.bold = True
    p.add_run(
        " к показателям G8/G9/G10. Ограничитель «x_raw × 2» — anti-washing защита: формальный "
        "декларативный уровень не может превратиться в уровень сертификации простым присвоением. "
        "Если x_raw = 0, бонус не начисляется (0 × 2 = 0): компания, не делающая ничего по стандарту, "
        "не получает амплификации за санкционный статус."
    )

    # Сначала — применённые бонусы
    if triggered:
        doc.add_heading("3.1. Применённые санкционные бонусы", level=2)
        rows = []
        for i in triggered:
            sb = i["sanctions_bonus"]
            rows.append([
                i["code"], i["name"],
                f"{sb['x_raw']:.2f}",
                f"+{sb['bonus_amount']:.2f}",
                f"{sb['x_final']:.2f}",
                sb["limit_used"],
                sb["trigger_reason"],
            ])
        add_multicol_table(doc,
            ["Код", "Показатель", "x без бонуса", "Бонус", "x с бонусом", "Ограничитель", "Основание"],
            rows, col_widths_cm=[1.2, 3.5, 2, 1.3, 2, 3.5, 3.9])

    # Затем — заявленные, но не применённые (для академической прозрачности)
    if not_triggered:
        doc.add_heading("3.2. Санкционные флаги без начисления бонуса", level=2)
        p = doc.add_paragraph()
        p.add_run(
            "По следующим показателям санкционный барьер заявлен, но бонус не начислен — потому что "
            "у компании нет раскрытой фактической практики, которую можно было бы амплифицировать. "
            "Это методологически корректный результат (anti-washing защита): причина низкой оценки — "
            "отсутствие раскрытия, а не санкционные ограничения."
        ).italic = True
        rows = []
        for i in not_triggered:
            sb = i["sanctions_bonus"]
            rows.append([
                i["code"], i["name"],
                f"{sb['x_raw']:.2f}",
                "0,00 (не начислен)",
                sb.get("non_application_reason", "x_raw = 0, бонус не применяется"),
            ])
        add_multicol_table(doc,
            ["Код", "Показатель", "x_raw", "Бонус", "Причина непримeнения"],
            rows, col_widths_cm=[1.2, 3.5, 1.5, 2.5, 8.7])


def write_risk_matrix(doc, data):
    doc.add_heading("4. Матрица отраслевых рисков (контур раскрытия)", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Для каждого из 10 отраслевых рисков золотодобычи проверяются 4 контура раскрытия: "
        "(1) внешний документ — попадает ли компания под требования закона или международного стандарта; "
        "(2) ВНД — принят ли внутренний нормативный документ; (3) Меры — описаны ли операционные действия; "
        "(4) КПЭ — раскрыт ли количественный показатель. Статус клеток привязан к фактическим x "
        "показателей 32-индикаторной модели (привязка — в последней колонке)."
    ).italic = True

    r = data["ratings"][-1] if data.get("dynamic_analysis") else data["ratings"][0]
    ind_map = {i["code"]: i for i in r["indicators"]}

    def status(bound_str):
        codes = [c.strip() for c in bound_str.split(",")]
        doc_flag = vnd_flag = meas_flag = kpi_flag = False
        for c in codes:
            i = ind_map.get(c)
            if not i or not i.get("applicable", True):
                continue
            x = i.get("x") or 0
            if x > 0: doc_flag = True
            if x >= 0.25: vnd_flag = True
            if x >= 0.50: meas_flag = True
            if x >= 0.75: kpi_flag = True
        return doc_flag, vnd_flag, meas_flag, kpi_flag

    rows, fills = [], []
    for risk_name, bound in INDUSTRY_RISKS:
        d, v, m, k = status(bound)
        rows.append([risk_name, "+" if d else "−", "+" if v else "−",
                    "+" if m else "−", "+" if k else "−", bound])
        fills.append([None,
                     "C8E6C9" if d else "FFCDD2",
                     "C8E6C9" if v else "FFCDD2",
                     "C8E6C9" if m else "FFCDD2",
                     "C8E6C9" if k else "FFCDD2",
                     None])
    add_multicol_table(doc,
        ["Отраслевой риск", "Документ", "ВНД", "Меры", "КПЭ", "Привязка"],
        rows, col_widths_cm=[5.8, 1.5, 1.3, 1.3, 1.3, 5.2],
        cell_fills=fills)
    p = doc.add_paragraph()
    p.add_run(
        "Интерпретация: все четыре «+» по риску означают полный контур раскрытия по АКРА-модели. "
        "«−» — недораскрытие соответствующего контура. Зелёная заливка — контур раскрыт, "
        "розовая — не раскрыт."
    ).italic = True


def write_block_details(doc, data):
    ratings_sorted = sorted(data["ratings"], key=lambda r: r["year"])
    primary = ratings_sorted[-1]
    section_num = 5
    for block in ["E", "S", "G"]:
        doc.add_heading(f"{section_num}. Блок {block} — {BLOCK_NAMES[block]}: декомпозиция", level=1)
        if len(ratings_sorted) > 1:
            doc.add_paragraph(
                f"Ниже детализация за {primary['year']} год (последний в динамике). "
                "Динамика по годам — в разделе «Динамический анализ»."
            )
        block_inds = [ind for ind in primary["indicators"] if ind["code"].startswith(block)]
        block_score = primary["block_scores"][block]
        doc.add_paragraph(
            f"Вес блока: 333 баллов. Компания получила: {block_score:.1f} "
            f"({100 * block_score / 333:.0f}% от максимума)."
        )

        # Сводная таблица блока с PDF-ссылками
        rows = []
        for ind in block_inds:
            if not ind.get("applicable", True):
                rows.append([ind["code"], ind["name"], "N/A", "N/A", "—"])
                continue
            pdf = ind.get("pdf_page_ref") or "—"
            if (ind.get("x") or 0) == 0:
                pdf = "— (нет данных)"
            rows.append([
                ind["code"],
                ind["name"],
                f"{ind['x']:.2f}" if ind["x"] is not None else "—",
                f"{ind['score']:.1f} / {ind['weight']}",
                pdf
            ])
        add_multicol_table(doc,
            ["Код", "Показатель", "x", "Балл", "Источник в PDF"],
            rows, col_widths_cm=[1.2, 6.5, 1.5, 2.5, 3.7])

        # Детальные подразделы
        for ind in block_inds:
            doc.add_heading(f"{ind['code']}. {ind['name']} (вес: {ind['weight']})", level=2)
            if not ind.get("applicable", True):
                p = doc.add_paragraph()
                p.add_run("⚠ Не применимо: ").bold = True
                p.add_run(ind.get("explanation", ""))
                continue
            lines = [
                ("x (нормированное значение)", f"{ind['x']:.4f}" if ind["x"] is not None else "—"),
                ("Балл", f"{ind['score']:.2f} / {ind['weight']}"),
            ]
            if ind.get("pdf_page_ref"):
                lines.append(("Источник в отчёте", ind["pdf_page_ref"]))
            if ind.get("raw_quote"):
                lines.append(("Цитата из отчёта", f"«{ind['raw_quote']}»"))
            add_kv_table(doc, lines)
            p = doc.add_paragraph()
            p.add_run("Обоснование: ").bold = True
            p.add_run(ind.get("explanation", ""))

            if (ind.get("sanctions_bonus") or {}).get("triggered"):
                sb = ind["sanctions_bonus"]
                p = doc.add_paragraph()
                p.add_run("🛡 Санкционный бонус применён: ").bold = True
                p.add_run(
                    f"x_raw={sb['x_raw']:.2f} → x_final={sb['x_final']:.2f} "
                    f"(ограничитель: {sb['limit_used']}). Основание: {sb['trigger_reason']}."
                )
            if ind.get("caveats_applied"):
                p = doc.add_paragraph()
                p.add_run("Применённые понижающие коэффициенты: ").bold = True
                p.add_run(", ".join(ind["caveats_applied"]))
        section_num += 1
    return section_num


def write_dynamic_analysis(doc, data, section_num):
    if "dynamic_analysis" not in data:
        return section_num
    da = data["dynamic_analysis"]
    doc.add_heading(f"{section_num}. Динамический анализ", level=1)
    doc.add_heading(f"{section_num}.1. История интегрального балла", level=2)
    add_multicol_table(doc, ["Год", "Балл", "Рейтинг"],
                       [[h["year"], f"{h['score']:.2f}", h["rating"]] for h in da["rating_history"]])
    doc.add_heading(f"{section_num}.2. ΔScore year-over-year", level=2)
    add_multicol_table(doc, ["Год", "Предыдущий год", "ΔScore"],
                       [[d["year"], d["prev_year"], f"{d['delta_score']:+.2f}"]
                        for d in da["year_over_year_deltas"]])
    if da.get("cagr") is not None:
        p = doc.add_paragraph()
        p.add_run(f"CAGR: {da['cagr']*100:+.2f}%").bold = True
    doc.add_heading(f"{section_num}.3. Декомпозиция по блокам", level=2)
    doc.add_paragraph(f"Период: {da['years'][0]} → {da['years'][-1]}")
    add_multicol_table(doc, ["Блок", "ΔScore"],
                       [[f"Блок {b} — {BLOCK_NAMES[b]}", f"{v:+.2f}"]
                        for b, v in da["block_deltas_first_to_last"].items()])
    doc.add_heading(f"{section_num}.4. Топ-5 драйверов роста", level=2)
    add_multicol_table(doc,
        ["Код", "Название", f"Балл {da['years'][0]}", f"Балл {da['years'][-1]}", "ΔScore"],
        [[r["code"], r["name"], f"{r['first_year_score']:.2f}",
          f"{r['last_year_score']:.2f}", f"{r['delta_score']:+.2f}"] for r in da["top5_risers"]])
    doc.add_heading(f"{section_num}.5. Топ-5 драйверов снижения", level=2)
    add_multicol_table(doc,
        ["Код", "Название", f"Балл {da['years'][0]}", f"Балл {da['years'][-1]}", "ΔScore"],
        [[r["code"], r["name"], f"{r['first_year_score']:.2f}",
          f"{r['last_year_score']:.2f}", f"{r['delta_score']:+.2f}"] for r in da["top5_fallers"]])
    return section_num + 1


def write_limitations(doc, data, section_num):
    doc.add_heading(f"{section_num}. Ограничения оценки", level=1)
    doc.add_paragraph(
        "Методология disclosure-based: оцениваются только раскрытые данные. Это накладывает ограничения:"
    )
    limitations = [
        "Скрытая информация (неотражённые инциденты, непубликуемые штрафы) не учитывается.",
        "Негативные события — только через количественные показатели и жёсткие потолки (S2, G4, S7).",
        "Нет forward-looking компоненты: оценивается фактическое состояние.",
        "Уровни зрелости 0–4 содержат элемент аналитического суждения.",
        "Санкционный бонус применяется только при подтверждённом ограничении и не освобождает от фактического соблюдения стандартов.",
        "Матрица отраслевых рисков имеет эвристическую привязку к показателям.",
    ]
    for l in limitations:
        doc.add_paragraph(f"• {l}", style='List Bullet')
    return section_num + 1


def write_conclusion(doc, data, section_num):
    doc.add_heading(f"{section_num}. Заключение", level=1)
    primary = sorted(data["ratings"], key=lambda r: r["year"])[-1]
    p = doc.add_paragraph()
    p.add_run(
        f"По результатам расчёта компания {primary['company']} за {primary['year']} год "
        f"получает интегральный ESG-балл {primary['total_score']:.1f} из 999, "
        f"что соответствует уровню {primary['rating']['label']} "
        f"({primary['rating']['interpretation'].lower()}). "
    )
    top_ind = sorted([i for i in primary["indicators"] if i.get("score")],
                     key=lambda x: -x["score"])[:3]
    p = doc.add_paragraph()
    p.add_run("Сильные стороны раскрытия: ").bold = True
    p.add_run(", ".join(f"{i['code']} ({i['name']}, {i['score']:.1f} б.)" for i in top_ind))
    p.add_run(".")

    not_disclosed = [(i["code"], i["name"], i["weight"]) for i in primary["indicators"]
                     if "not_disclosed" in i.get("caveats_applied", []) and i.get("applicable", True)]
    not_disclosed.sort(key=lambda x: -x[2])
    if not_disclosed:
        p = doc.add_paragraph()
        p.add_run("Ключевые зоны для улучшения раскрытия: ").bold = True
        p.add_run(", ".join(f"{c} ({n}, вес {w})" for c, n, w in not_disclosed[:5]))
        p.add_run(". См. раздел «Потенциал роста» для референсов по отраслевым лучшим практикам.")
    return section_num + 1


def write_appendix(doc, data):
    doc.add_page_break()
    doc.add_heading("Приложение. Детальная таблица всех 32 показателей", level=1)
    primary = sorted(data["ratings"], key=lambda r: r["year"])[-1]
    rows = []
    for ind in primary["indicators"]:
        src = ind.get("pdf_page_ref") or "—"
        if not ind.get("applicable", True):
            rows.append([ind["code"], ind["name"], ind["weight"], "N/A", "N/A", src])
        else:
            x_str = f"{ind['x']:.3f}" if ind["x"] is not None else "—"
            s_str = f"{ind['score']:.2f}" if ind["score"] is not None else "—"
            rows.append([ind["code"], ind["name"], ind["weight"], x_str, s_str, src])
    rows.append(["ИТОГО", "", 999, "", f"{primary['total_score']:.1f}", ""])
    add_multicol_table(doc,
        ["Код", "Показатель", "Вес", "x", "Балл", "Источник"],
        rows, col_widths_cm=[1.2, 7.0, 1.2, 1.5, 2.5, 3.0])


def _fix_zoom_attribute(docx_path):
    """Костыль: python-docx создаёт <w:zoom w:val="bestFit"/> без атрибута percent,
    который требует OOXML schema. Пере-пакуем файл с исправленным settings.xml."""
    import zipfile, shutil, tempfile, os
    tmpdir = tempfile.mkdtemp()
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            z.extractall(tmpdir)
        settings_path = os.path.join(tmpdir, 'word', 'settings.xml')
        if os.path.exists(settings_path):
            with open(settings_path, 'r', encoding='utf-8') as f:
                content = f.read()
            new_content = content.replace(
                '<w:zoom w:val="bestFit"/>',
                '<w:zoom w:val="bestFit" w:percent="100"/>'
            )
            if new_content != content:
                with open(settings_path, 'w', encoding='utf-8') as f:
                    f.write(new_content)
                # пере-пакуем
                tmp_zip = docx_path + '.tmp'
                with zipfile.ZipFile(tmp_zip, 'w', zipfile.ZIP_DEFLATED) as zout:
                    for root, _, files in os.walk(tmpdir):
                        for fname in files:
                            fpath = os.path.join(root, fname)
                            arcname = os.path.relpath(fpath, tmpdir)
                            zout.write(fpath, arcname)
                shutil.move(tmp_zip, docx_path)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("result_json")
    parser.add_argument("output_docx")
    parser.add_argument("--meta", default=None)
    args = parser.parse_args()

    with open(args.result_json, "r", encoding="utf-8") as f:
        data = json.load(f)
    meta = None
    if args.meta:
        with open(args.meta, "r", encoding="utf-8") as f:
            meta = json.load(f)

    doc = Document()
    set_style_defaults(doc)
    write_title(doc, data)
    write_object(doc, data)
    write_summary(doc, data)
    if meta:
        write_meta_indicators(doc, data, meta)
    write_best_worst(doc, data)
    write_sanctions_adjustment(doc, data)
    write_risk_matrix(doc, data)
    sec = write_block_details(doc, data)
    sec = write_dynamic_analysis(doc, data, sec)
    sec = write_limitations(doc, data, sec)
    sec = write_conclusion(doc, data, sec)
    write_appendix(doc, data)
    doc.save(args.output_docx)
    _fix_zoom_attribute(args.output_docx)
    print(f"Отчёт сохранён: {args.output_docx}")


if __name__ == "__main__":
    main()
